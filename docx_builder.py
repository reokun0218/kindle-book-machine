# docx_builder.py — builds a beautifully designed Word .docx file from book_data
import sys
sys.stdout.reconfigure(encoding="utf-8", errors="replace")
from pathlib import Path
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT_DIR = Path("output")

NUM_KANJI = {
    1: "一", 2: "二", 3: "三", 4: "四", 5: "五",
    6: "六", 7: "七", 8: "八", 9: "九", 10: "十",
    11: "十一", 12: "十二", 13: "十三", 14: "十四", 15: "十五"
}
NUM_WORDS_EN = {
    1: "One", 2: "Two", 3: "Three", 4: "Four", 5: "Five",
    6: "Six", 7: "Seven", 8: "Eight", 9: "Nine", 10: "Ten",
    11: "Eleven", 12: "Twelve", 13: "Thirteen", 14: "Fourteen", 15: "Fifteen"
}


def _set_font(run, name="MS Gothic", size=11, bold=False, italic=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)
    r = run._r
    rPr = r.get_or_add_rPr()
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:eastAsia"), name)
    rPr.insert(0, rFonts)


def _add_paragraph(doc, text="", align=WD_ALIGN_PARAGRAPH.LEFT,
                   size=11, bold=False, italic=False, color=None,
                   space_before=0, space_after=6, line_spacing=18,
                   indent=False, center=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    p.paragraph_format.line_spacing = Pt(line_spacing)
    if indent:
        p.paragraph_format.first_line_indent = Inches(0.25)
    if text:
        run = p.add_run(text)
        _set_font(run, size=size, bold=bold, italic=italic, color=color)
    return p


def _add_decorative_line(doc, color_rgb=(139, 92, 246)):
    """Add a thin colored horizontal rule."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    hex_color = "{:02X}{:02X}{:02X}".format(*color_rgb)
    bottom.set(qn("w:color"), hex_color)
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def _add_chapter_heading(doc, num_label, title, is_japanese=True):
    """Add a beautifully styled chapter heading with decorative line."""
    # Chapter number label (e.g. 第一章 or Chapter One)
    p_num = doc.add_paragraph()
    p_num.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_num.paragraph_format.space_before = Pt(12)
    p_num.paragraph_format.space_after  = Pt(2)
    run_num = p_num.add_run(num_label)
    _set_font(run_num, size=10, color=(139, 92, 246), bold=True)

    # Title
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_title.paragraph_format.space_before = Pt(2)
    p_title.paragraph_format.space_after  = Pt(6)
    run_title = p_title.add_run(title)
    _set_font(run_title, size=18, bold=True)

    # Decorative line under heading
    _add_decorative_line(doc)

    # Space after line
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def _add_body_text(doc, text, indent=True):
    """Add multi-paragraph body text."""
    if not text:
        return
    paragraphs = [p.strip() for p in text.split("\n\n") if p.strip()]
    for i, para_text in enumerate(paragraphs):
        # Handle markdown-style bold
        if "**" in para_text:
            p = doc.add_paragraph()
            p.paragraph_format.space_after  = Pt(6)
            p.paragraph_format.line_spacing = Pt(20)
            if indent and i > 0:
                p.paragraph_format.first_line_indent = Inches(0.25)
            # Split on **bold** markers
            parts = para_text.split("**")
            for j, part in enumerate(parts):
                if not part:
                    continue
                part = part.replace("\n", " ")
                run = p.add_run(part)
                _set_font(run, size=11, bold=(j % 2 == 1))
        elif para_text.startswith("---") or para_text.startswith("—-"):
            _add_decorative_line(doc, color_rgb=(200, 200, 200))
        else:
            p = doc.add_paragraph()
            p.paragraph_format.space_after  = Pt(6)
            p.paragraph_format.line_spacing = Pt(20)
            if indent and i > 0:
                p.paragraph_format.first_line_indent = Inches(0.25)
            run = p.add_run(para_text.replace("\n", " "))
            _set_font(run, size=11)


def _add_upsell_page(doc, book_data, is_japanese):
    """Add a beautifully designed bonus/CTA page at the end."""
    has_any = any([
        book_data.get("bonus_title"),
        book_data.get("line_url"),
        book_data.get("consultation_url"),
        book_data.get("upsell_product"),
        book_data.get("lead_magnet"),
    ])
    if not has_any:
        return

    doc.add_page_break()

    label = "🎁 読者への特典・次のステップ" if is_japanese else "🎁 Bonus & Next Steps"
    _add_paragraph(doc, label, size=18, bold=True, center=True, space_before=20, space_after=4)
    _add_decorative_line(doc, color_rgb=(236, 72, 153))
    _add_paragraph(doc, "", space_after=8)

    if book_data.get("bonus_title"):
        _add_paragraph(doc, "【 無料特典 】" if is_japanese else "[ FREE BONUS ]",
                       size=10, color=(139, 92, 246), bold=True, space_after=2)
        _add_paragraph(doc, book_data["bonus_title"], size=14, bold=True, space_after=4)
        if book_data.get("bonus_description"):
            _add_paragraph(doc, book_data["bonus_description"], size=11, space_after=12)

    if book_data.get("lead_magnet"):
        icon = "🧲 " if is_japanese else "🧲 "
        _add_paragraph(doc, icon + ("無料プレゼント" if is_japanese else "Free Gift"),
                       size=10, color=(34, 211, 238), bold=True, space_after=2)
        _add_paragraph(doc, book_data["lead_magnet"], size=12, bold=True, space_after=10)

    if book_data.get("line_url"):
        _add_paragraph(doc, "📲 " + ("公式LINEに登録する" if is_japanese else "Join Official LINE"),
                       size=12, bold=True, color=(34, 197, 94), space_after=2)
        _add_paragraph(doc, book_data["line_url"], size=11, italic=True, space_after=10)

    if book_data.get("consultation_url"):
        _add_paragraph(doc, "🤝 " + ("個別相談・セッションを申し込む" if is_japanese else "Book a Consultation"),
                       size=12, bold=True, color=(139, 92, 246), space_after=2)
        _add_paragraph(doc, book_data["consultation_url"], size=11, italic=True, space_after=10)

    if book_data.get("upsell_product"):
        _add_paragraph(doc, "💎 " + ("次のステップ" if is_japanese else "Next Step"),
                       size=10, color=(236, 72, 153), bold=True, space_after=2)
        _add_paragraph(doc, book_data["upsell_product"], size=13, bold=True, space_after=10)

    _add_decorative_line(doc, color_rgb=(139, 92, 246))
    closing = "この本を最後まで読んでくださり、ありがとうございました。\nあなたの人生が、今日から少しずつ変わっていくことを願っています。" \
              if is_japanese else \
              "Thank you for reading. I believe in your transformation."
    _add_paragraph(doc, closing, size=11, italic=True, space_before=14, space_after=0, center=True)

    author = book_data.get("author_display_name", "")
    if author:
        _add_paragraph(doc, f"— {author}", size=11, bold=True, space_before=8, center=True)


def create_docx(book_data, output_filename):
    """Build and save the DOCX file. Returns the output file path."""
    OUTPUT_DIR.mkdir(exist_ok=True)

    language    = book_data.get("language", "english")
    is_japanese = language == "japanese"
    chapters    = book_data.get("chapters", [])

    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Inches(1.1)
        section.bottom_margin = Inches(1.1)
        section.left_margin   = Inches(1.3)
        section.right_margin  = Inches(1.3)

    # ── TITLE PAGE ────────────────────────────────────────────────────────
    _add_paragraph(doc, "", space_before=0, space_after=60)

    # Decorative top line
    _add_decorative_line(doc, color_rgb=(139, 92, 246))
    _add_paragraph(doc, "", space_after=16)

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(8)
    title_p.paragraph_format.space_after  = Pt(8)
    title_run = title_p.add_run(book_data.get("title", ""))
    _set_font(title_run, size=26, bold=True)

    subtitle = book_data.get("subtitle", "")
    if subtitle:
        sub_p = doc.add_paragraph()
        sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sub_p.paragraph_format.space_before = Pt(6)
        sub_p.paragraph_format.space_after  = Pt(6)
        sub_run = sub_p.add_run(subtitle)
        _set_font(sub_run, size=13, italic=True, color=(139, 92, 246))

    _add_paragraph(doc, "", space_after=32)
    _add_decorative_line(doc, color_rgb=(236, 72, 153))
    _add_paragraph(doc, "", space_after=40)

    author_p = doc.add_paragraph()
    author_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    author_run = author_p.add_run(book_data.get("author_display_name", ""))
    _set_font(author_run, size=14, bold=True)

    year_p = doc.add_paragraph()
    year_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    year_p.paragraph_format.space_before = Pt(6)
    year_run = year_p.add_run(str(datetime.now().year))
    _set_font(year_run, size=11, color=(150, 150, 150))

    credit_p = doc.add_paragraph()
    credit_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    credit_p.paragraph_format.space_before = Pt(48)
    credit_run = credit_p.add_run("CREATED BY 渡邉有優美 × Claude AI")
    _set_font(credit_run, size=9, color=(120, 100, 160))

    doc.add_page_break()

    # ── TABLE OF CONTENTS ─────────────────────────────────────────────────
    toc_label        = "目次" if is_japanese else "Table of Contents"
    intro_label      = "はじめに" if is_japanese else "Introduction"
    conclusion_label = "おわりに" if is_japanese else "Conclusion"

    _add_paragraph(doc, toc_label, size=18, bold=True, space_before=8, space_after=4)
    _add_decorative_line(doc)
    _add_paragraph(doc, "", space_after=8)

    toc_items = [("→ " + intro_label, False)]
    for ch in chapters:
        num = ch["number"]
        if is_japanese:
            label = f"第{NUM_KANJI.get(num, str(num))}章　{ch['title']}"
        else:
            label = f"Chapter {NUM_WORDS_EN.get(num, str(num))}: {ch['title']}"
        toc_items.append(("→ " + label, False))
    toc_items.append(("→ " + conclusion_label, False))

    if any([book_data.get("bonus_title"), book_data.get("line_url"),
            book_data.get("consultation_url"), book_data.get("upsell_product")]):
        bonus_label = "🎁 特典・次のステップ" if is_japanese else "🎁 Bonus & Next Steps"
        toc_items.append(("→ " + bonus_label, True))

    for item, is_bonus in toc_items:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after  = Pt(2)
        run = p.add_run(item)
        _set_font(run, size=11, bold=is_bonus,
                  color=(236, 72, 153) if is_bonus else None)

    doc.add_page_break()

    # ── INTRODUCTION ──────────────────────────────────────────────────────
    _add_chapter_heading(doc, "―" if is_japanese else "―",
                         intro_label, is_japanese)
    _add_body_text(doc, book_data.get("introduction", ""))
    doc.add_page_break()

    # ── CHAPTERS ──────────────────────────────────────────────────────────
    for ch in chapters:
        num = ch["number"]
        if is_japanese:
            num_label = f"第{NUM_KANJI.get(num, str(num))}章"
        else:
            num_label = f"Chapter {NUM_WORDS_EN.get(num, str(num))}"
        _add_chapter_heading(doc, num_label, ch["title"], is_japanese)
        _add_body_text(doc, ch.get("content", ""))
        doc.add_page_break()

    # ── CONCLUSION ────────────────────────────────────────────────────────
    _add_chapter_heading(doc, "―" if is_japanese else "―",
                         conclusion_label, is_japanese)
    _add_body_text(doc, book_data.get("conclusion", ""))

    # ── BONUS / UPSELL PAGE ───────────────────────────────────────────────
    _add_upsell_page(doc, book_data, is_japanese)

    doc.add_page_break()

    # ── ABOUT AUTHOR ──────────────────────────────────────────────────────
    about_label = "著者について" if is_japanese else "About the Author"
    _add_chapter_heading(doc, "―", about_label, is_japanese)
    author_name = book_data.get("author_display_name", "[著者名]")
    if is_japanese:
        placeholder = f"{author_name} は、[職業・肩書き] として活躍中です。\n[著者の簡単な紹介をここに入れてください。SNSアカウント、ウェブサイトなども記載できます。]"
    else:
        placeholder = f"{author_name} is a [title/profession].\n[Add your 2-3 sentence bio here. Include your website and social media links.]"
    _add_body_text(doc, placeholder)

    # ── SAVE ──────────────────────────────────────────────────────────────
    safe_name = output_filename.replace(" ", "_")
    if not safe_name.endswith(".docx"):
        safe_name += ".docx"
    out_path = OUTPUT_DIR / safe_name
    doc.save(str(out_path))
    size_kb = out_path.stat().st_size // 1024
    print(f"✅ DOCX created: {out_path} ({size_kb}KB)")
    return str(out_path)
